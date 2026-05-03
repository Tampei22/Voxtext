"""Export an STTSession to TXT, DOCX, SRT, or PDF."""
from app_core.models import STTSession


def _fmt_time(seconds: float) -> str:
    """Format seconds as SRT timestamp HH:MM:SS,mmm."""
    h = int(seconds // 3600)
    m = int((seconds % 3600) // 60)
    s = int(seconds % 60)
    ms = int((seconds % 1) * 1000)
    return f"{h:02d}:{m:02d}:{s:02d},{ms:03d}"


def export_txt(session: STTSession, path: str) -> None:
    with open(path, "w", encoding="utf-8") as f:
        f.write(session.full_text)


def export_srt(session: STTSession, path: str) -> None:
    with open(path, "w", encoding="utf-8") as f:
        for i, phrase in enumerate(session.phrases, 1):
            f.write(f"{i}\n")
            f.write(f"{_fmt_time(phrase.start)} --> {_fmt_time(phrase.end)}\n")
            f.write(f"{phrase.text.strip()}\n\n")


def export_docx(session: STTSession, path: str) -> None:
    from docx import Document
    doc = Document()
    doc.add_heading(f"STT Session — {session.created_at_iso[:10]}", level=1)
    info = doc.add_paragraph()
    info.add_run(f"Language: {session.lang}   Engine: {session.engine}")
    doc.add_paragraph()
    for phrase in session.phrases:
        p = doc.add_paragraph()
        ts_run = p.add_run(f"[{_fmt_time(phrase.start)}]  ")
        ts_run.bold = True
        p.add_run(phrase.text.strip())
    doc.save(path)


def export_pdf(session: STTSession, path: str) -> None:
    try:
        from fpdf import FPDF
        _pdf_fpdf2(session, path, FPDF)
        return
    except ImportError:
        pass
    try:
        import reportlab  # noqa: F401
        _pdf_reportlab(session, path)
        return
    except ImportError:
        pass
    raise ImportError(
        "PDF export requires 'fpdf2' or 'reportlab'.\n"
        "Install with:  pip install fpdf2"
    )


def _pdf_fpdf2(session: STTSession, path: str, FPDF) -> None:
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_font("Helvetica", "B", 13)
    pdf.cell(0, 10, f"STT Session  {session.created_at_iso[:10]}",
             new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", size=10)
    pdf.cell(0, 8, f"Language: {session.lang}   Engine: {session.engine}",
             new_x="LMARGIN", new_y="NEXT")
    pdf.ln(4)
    pdf.set_font("Helvetica", size=11)
    for phrase in session.phrases:
        ts = f"[{_fmt_time(phrase.start)}]"
        safe = phrase.text.strip().encode("latin-1", errors="replace").decode("latin-1")
        pdf.multi_cell(0, 7, f"{ts}  {safe}")
    pdf.output(path)


def _pdf_reportlab(session: STTSession, path: str) -> None:
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm
    c = canvas.Canvas(path, pagesize=A4)
    page_w, page_h = A4
    y = page_h - 2 * cm
    c.setFont("Helvetica-Bold", 13)
    c.drawString(2 * cm, y, f"STT Session — {session.created_at_iso[:10]}")
    y -= 0.7 * cm
    c.setFont("Helvetica", 10)
    c.drawString(2 * cm, y, f"Language: {session.lang}   Engine: {session.engine}")
    y -= 1.0 * cm
    c.setFont("Helvetica", 11)
    for phrase in session.phrases:
        ts = f"[{_fmt_time(phrase.start)}]"
        try:
            safe = phrase.text.strip().encode("latin-1", errors="replace").decode("latin-1")
        except Exception:
            safe = ""
        line = f"{ts}  {safe}"
        if y < 2 * cm:
            c.showPage()
            y = page_h - 2 * cm
            c.setFont("Helvetica", 11)
        c.drawString(2 * cm, y, line[:110])
        y -= 0.65 * cm
    c.save()
